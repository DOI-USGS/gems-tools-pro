"""
Translates GeMS-style DMU table into a USGS-formatted Microsoft MS Word .docx file.
Table may be in any format recognized by ArcGIS Pro, eg., table view in a map, file gdb, 
geopackage, dbf, csv, Excel sheet.

Arguments
    dmu_table (str) : Path to description of map units table.
    out_dir (str) : Path to output folder.
    out_name (str) : Name of the output docx file. Extension appended if not included.
    calc_style (boolean) : Should paragraph styles be determined programmatically (True)
        or should the value in ParagraphStyle (False) be used? True by default.
    use_label (boolean) : Should the Label field value be used for the unit abbreviations (True)
        or should the MapUnit field be used (False)? False by default
    format (boolean) : Should in-line HTML and/or ArcGIS text formatting tags in the Description 
        field be converted to Word run properties? True by default.
    is_lmu (boolean) : Make a 'List of Map Units' with no descriptions (True) instead of 
        a 'Description of Map Units' (False)? False by default.
    open_doc (boolean) : Open the resulting docx file in Word for immediate viewing and editing? 
        True by default.

Dependencies
    docx (https://python-docx.readthedocs.io/en/latest/) included with toolbox in folder
        Scripts\docx
    BeautifulSoup - installed in default ArcGIS Pro miniconda environment
"""

import sys
import os
import arcpy
from pathlib import Path
import GeMS_utilityFunctions as guf
import docx
import bs4
from bs4 import BeautifulSoup

versionString = "GeMS_DMUtoDocx.py, version of 5/8/2024"
rawurl = "https://raw.githubusercontent.com/DOI-USGS/gems-tools-pro/master/Scripts/GeMS_DMUToDocx.py"
guf.checkVersion(versionString, rawurl, "gems-tools-pro")

# value[1] of style_dict can be used in Word docs based on USGS MapManuscript_v3-1_06-22.dotx
# older template MapManuscript_v1-0_04-11.dotx not guaranteed to work.
# multiple key values included here in an attempt to allow variations on the style name
style_dict = {
    "dmu - list bullet": ["unit text", "DMU - List Bullet"],
    "dmu headnote - 1 line": ["headnote", "DMU Headnote - 1 Line"],
    "dmu headnote - more than 1 line": ["headnote", "DMU Headnote - More Than 1 Line"],
    "dmu headnote paragraph": ["headnote", "DMU Headnote Paragraph"],
    "dmuheadnote": ["headnote", "DMU Headnote - More Than 1 Line"],
    "dmu noindent": ["unit text", "DMU NoIndent"],
    "dmu paragraph": ["unit text", "DMU Paragraph"],
    "dmu quotation": ["unit text", "DMU Quotation"],
    "dmu unit 1 (1st after heading)": ["unit", "DMU Unit 1 (1st after heading)"],
    "dmu unit 1": ["unit", "DMU Unit 1"],
    "dmuunit1": ["unit", "DMU Unit 1"],
    "dmu unit 2": ["unit", "DMU Unit 2"],
    "dmuunit2": ["unit", "DMU Unit 2"],
    "dmu unit 3": ["unit", "DMU Unit 3"],
    "dmuunit3": ["unit", "DMU Unit 3"],
    "dmu unit 4": ["unit", "DMU Unit 4"],
    "dmuunit4": ["unit", "DMU Unit 4"],
    "dmu unit 5": ["unit", "DMU Unit 5"],
    "dmuunit5": ["unit", "DMU Unit 5"],
    "dmu unit label (type style)": ["label", "DMU Unit Label (type style)"],
    "dmu unit name/age (type style)": ["age", "DMU Unit Name/Age (type style)"],
    "dmu-heading1": ["heading", "DMU-Heading1"],
    "dmuheading1": ["heading", "DMU-Heading1"],
    "dmu-heading2": ["heading", "DMU-Heading2"],
    "dmuheading2": ["heading", "DMU-Heading2"],
    "dmu-heading3": ["heading", "DMU-Heading3"],
    "dmuheading3": ["heading", "DMU-Heading3"],
    "dmu-heading4": ["heading", "DMU-Heading4"],
    "dmuheading4": ["heading", "DMU-Heading4"],
    "dmu-heading5": ["heading", "DMU-Heading5"],
    "dmuheading5": ["heading", "DMU-Heading5"],
    "run-inhead": ["headnote", "Run-inHead"],
}


def strip_string(a):
    """Strip leading and trailing spaces from a string but check first
    that the string variable is not None."""
    if a:
        return str(a).strip()
    else:
        return a


def apply_formatting(el, run):
    """Translates a few ArcGIS text formatting and HTML tags into
    docx run properties"""
    # ArcGIS text formatting tag FNT with size and style
    if str(el.name).lower() == "fnt":
        if el.get("name"):
            run.font.name = el.get("name")
        if el.get("size"):
            size = int(el.get("size"))
            run.font.size = docx.shared.Pt(size)
        if el.get("style"):
            if el.get("style").lower() == "italic":
                run.font.italic = True
            if el.get("style").lower() == "regular":
                run.font.italic = False
        # assume any weight is equivalent to 'bold'
        if el.get("wght"):
            run.font.bold = True

        return

    # HTML <span with in-line css in style attribute
    # <span style='font-family: FGDCGeoAge; font-weight: bold; font-style: italic;'
    if str(el.name).lower() == "span":
        if el.get("style"):
            style = el.get("style").strip("'").strip('"')
            style_parse = style.split(";")
            for el in style_parse:
                if "font-family" in el:
                    font = el.split("font-family:")[1].strip()
                    run.font.name = font
                if "font-style" in el:
                    font_style = el.split("font-style:")[1].strip()
                    if font_style.lower() == "italic":
                        run.font.italic = True
                if "font-weight" in el:
                    weight = el.split("font-weight:")[1].strip()
                    if weight == "bold":
                        run.font.bold = True
                    if weight == "normal":
                        run.font.bold = False
        return

    # standalone ArcGIS and HTML tags
    if str(el.name).lower() in ("bol", "b", "strong"):
        run.font.bold = True
        return

    if str(el.name).lower() == "sup":
        run.font.superscript = True
        return

    if str(el.name).lower() in ("em", "i", "ita"):
        run.font.italic = True
        return

    if str(el.name).lower() == "sub":
        run.font.subscript = True
        return

    return


def determine_style(para_type, i, rows):
    """Determines the specific Word doc style to apply to the paragraph based
    on paragraph type and properties of the previously inserted row or rows"""
    # rows = [hkey, unit, name, age, description, paragraph style]
    # the first item in the list has already been set as Description or
    # List of Map Units and this will always be DMU-Heading1
    if i == 0:
        return "dmu-heading1"

    # headnotes do not need to be ranked
    if para_type == "headnote":
        if rows[i][4].startswith("["):
            if len(rows[i][4]) <= 91:
                return "dmu headnote - 1 line"
            else:
                return "dmu headnote - more than 1 line"
        else:
            return "dmu headnote paragraph"

    # continue for headings and units
    # hkey of this row
    i_hkey = rows[i][0]

    # properties of previous row
    p_hkey = rows[i - 1][0]
    p_para_style = rows[i - 1][5]
    p_para_type = style_dict[p_para_style.lower()][0]

    # unit after headnote or heading
    if para_type == "unit":
        if p_para_type == "heading":
            return "dmu unit 1 (1st after heading)"
        if p_para_type == "headnote":
            return "dmu unit 1"

    # now go backwards through list looking for the last paragraph of the same type
    for row in reversed(rows[0:i]):
        if style_dict[row[5].lower()][0] == para_type:
            # special case of discovering Heading 2 after Heading 1
            if row[5].lower() == "dmu-heading1":
                return "dmu-heading2"

            # same type and same hkey length means these are siblings
            # use the same style as the older sibling
            if len(row[0]) == len(i_hkey):
                if row[5].lower() == "dmu unit 1 (1st after heading)":
                    return "dmu unit 1"
                else:
                    return row[5]

            # same type but current hkey is longer, the current is a child
            # of the last one found, increment the style rank by one
            if len(i_hkey) > len(row[0]):
                r_para_style = row[5]
                if r_para_style == "dmu unit 1 (1st after heading)":
                    return "dmu unit 2"
                else:
                    r_rank_i = int(r_para_style[-1])
                    i_unit_rank = r_rank_i + 1
                    return f"{r_para_style[0:-1]}{i_unit_rank}"

        # special situation where a unit has a hierarchy key of the same length as a previous heading.
        # eg. BEDROCK UNITS is followed by a number of real bedrock units but those are followed by something
        # like 'water', 'unmapped', or 'ice, that is, unclassified units at the end of the list.
        if len(i_hkey) == len(row[0]) and para_type == "unit" and "heading" in row[5]:
            arcpy.AddWarning(
                f"The unit '{rows[i][1]}' (HierarchyKey: {i_hkey}) and other units that may follow it, appear to be out of order.\n"
                f"Review its HierarchyKey or consider adding a heading above it to differentiate it from '{row[2]}'"
            )

            return "dmu unit 1"

    # finally
    arcpy.AddWarning(
        f"Could not determine a paragraph style for row with HierarchyKey {i_hkey}.\n"
        "Review the DMU table for inconsistent or missing values\n"
        "This paragraph will be styled with 'DMU Unit 1'"
    )
    return "dmu unit 1"


def determine_type(vals):
    """Determines the general paragraph type from table field values"""
    unit = vals[0]
    name = vals[1]
    age = vals[2]
    desc = vals[3]

    if unit:
        return "unit"
    if not unit and name:
        return "heading"
    if not unit and not name and desc:
        return "headnote"

    return None


def iterate_soup(paragraph, text, special=None):
    """Iterates through descendants of a bs4 soup and translates formatting tags into docx run properties.
    Either applies the properties here in this function or calls apply_formatting()"""
    soup = BeautifulSoup(text, "html.parser")

    i = 0
    for child in soup.descendants:
        if type(child) == bs4.element.NavigableString:
            run = paragraph.add_run()
            run.text = child.text

            if special == "unit":
                run.style = "DMU Unit Label (type style)"

            parents = [p for p in child.fetchParents() if not p.name == "[document]"]
            if parents:
                if parents[0].name in ("i", "em") and i == 0 and special == "headnote":
                    run.style = "Run-inHead"
                    i = i + 1
                    continue

                for parent in parents:
                    apply_formatting(parent, run)


def main(params):
    # PARAMETERS
    dmu_table = params[0]

    if not arcpy.Exists(dmu_table):
        arcpy.AddError("Could not find DescriptionOfMapUnits table")
        sys.exit()

    out_dir = Path(params[1])
    out_name = params[2]
    if not out_name.endswith(".docx"):
        out_name = f"{out_name}.docx"
    out_file = out_dir / out_name

    # do we calculate paragraph styles based on dmu attributes or use the value in ParagraphStyle?
    calc_style = False
    if len(params) > 3:
        calc_style = guf.eval_bool(params[3])

    # do we use the value in MapUnit or Label for the unit abbreviations in the docx?
    use_label = False
    if len(params) > 4:
        use_label = guf.eval_bool(params[4])

    # do we try to translate annotation text formatting into Word styles?
    format = False
    if len(params) > 5:
        format = guf.eval_bool(params[5])

    if use_label:
        mu = "Label"
    else:
        mu = "MapUnit"

    if calc_style:
        fields = ["HierarchyKey", mu, "Name", "Age", "Description"]
    else:
        fields = ["HierarchyKey", mu, "Name", "Age", "Description", "ParagraphStyle"]

    # are we making a DMU or an LMU?
    is_lmu = False
    if len(params) > 6:
        is_lmu = guf.eval_bool(params[6])

    open_doc = False
    if len(params) > 7:
        open_doc = guf.eval_bool(params[7])

    # START
    sqlclause = (None, "ORDER by HierarchyKey ASC")
    rows = [
        list(row)
        for row in arcpy.da.SearchCursor(dmu_table, fields, sql_clause=sqlclause)
    ]

    # strings might have leading or trailing spaces
    rows = [list(map(strip_string, row)) for row in rows]

    if is_lmu == True:
        head1 = "LIST OF MAP UNITS"
    else:
        head1 = "DESCRIPTION OF MAP UNITS"

    # add LMU or DMU title to rows
    if rows[0][2]:
        if not rows[0][2].lower().strip() == head1.lower():
            rows.insert(0, ["000", None, head1, None, None])
    else:
        rows.insert(0, ["000", None, head1, None, None])

    if not calc_style:
        rows[0].append("dmu-heading1")

    # rows = [hkey, unit, name, age, description]
    # doc_rows = [hkey, unit, name, age, description, paragraph style]
    for i, row in enumerate(rows):
        para_type = determine_type(row[1:5])
        if not calc_style:
            para_style = row[5]
        else:
            para_style = determine_style(para_type, i, rows)
        row.append(para_style)

    # remove headnotes if this is a list of map units
    for row in rows:
        if is_lmu and style_dict[row[5].lower()][0] == "headnote":
            rows.remove(row)

    # rows = [hkey, unit, name, age, description, paragraph style]
    scripts = Path.cwd()
    toolbox = scripts.parent
    resources = toolbox / "Resources"
    template = resources / "DMU_template.docx"

    document = docx.Document(template)
    document._body.clear_content()
    arcpy.AddMessage("Evaluating table")
    for p in rows:
        contents = ", ".join([n for n in p[1:3] if n])
        if p[2]:
            arcpy.AddMessage(f"  {p[2]}")
        else:
            arcpy.AddMessage(f"    -headnote text")

        if calc_style == False:
            if not p[5]:
                arcpy.AddError(
                    "Null value found in ParagraphStyle. Add value or choose 'Calculate paragraph style' and try again."
                )
                sys.exit()

        if style_dict[p[5].lower()][0] == "heading":
            document.add_paragraph(p[2], style_dict[p[5].lower()][1])

        if style_dict[p[5].lower()][0] == "headnote":
            for para in p[4].splitlines():
                headnote = document.add_paragraph(style=style_dict[p[5].lower()][1])
                if format:
                    iterate_soup(headnote, para, "headnote")
                else:
                    headnote.text = para

        if style_dict[p[5].lower()][0] == "unit":
            unit = document.add_paragraph(style=style_dict[p[5].lower()][1])
            if use_label:
                iterate_soup(unit, p[1], "unit")
            else:
                unit.add_run(f"{p[1]}", "DMU Unit Label (type style)")

            unit.add_run(f"\t{p[2]} ({p[3]})", "DMU Unit Name/Age (type style)")

            if not is_lmu:
                paras = None
                if p[4]:
                    paras = p[4].splitlines()

                if paras:
                    # add the first paragraph with style based on unit rank
                    # prepend an em-dash
                    if format:
                        iterate_soup(unit, f"—{paras[0]}")
                    else:
                        unit.add_run(f"—{paras[0]}")

                    # add the rest of the paragraphs with style DMU Paragraph
                    for para in [n for n in paras[1:] if n]:
                        new_p = document.add_paragraph(style="DMU Paragraph")
                        if format:
                            iterate_soup(new_p, para)
                        else:
                            new_p.add_run(para)

    arcpy.AddMessage(f"Saving {out_file}")
    try:
        document.save(out_file)
        del document
    except IOError:
        arcpy.AddError(
            f"Cannot save changes to {out_file}. If it is open, close it and try again."
        )
        sys.exit()

    if open_doc:
        os.startfile(out_file, "edit")


if __name__ == "__main__":
    main(sys.argv[1:])
